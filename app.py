import asyncio
import base64
import hmac
import hashlib
import json
import logging
import os
import re
import time
from contextlib import asynccontextmanager
from datetime import datetime, timedelta
from html import escape
from typing import Dict, List, Optional

import httpx
try:
    from ebooklib import epub
except ImportError:
    epub = None

try:
    from docx import Document
except ImportError:
    Document = None

from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, Response, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sqlalchemy import case, delete, select, func, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from sse_starlette.sse import EventSourceResponse

import config
import settings_manager
from database import init_db, migrate_db, get_db, async_session
from models import Novel, Chapter, ExpandTask
from parser import parse_novel_from_bytes
from ai_service import (
    AIRefusalError,
    ExpansionIntegrityError,
    expand_chapter_one_pass,
    chat_completion,
    stream_rewrite_paragraph,
    split_into_paragraphs,
    merge_paragraphs,
    generate_chapter_summary,
    build_local_chapter_summary,
    normalize_output_text,
    get_prompt_settings,
    update_prompt_settings,
    reset_prompt_settings,
    get_default_rewrite_instruction,
)

# ========== 日志配置 ==========
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

# ========== 全局状态 ==========
# SSE 订阅者: novel_id -> list of asyncio.Queue
sse_subscribers: Dict[int, List[asyncio.Queue]] = {}

# 取消信号: task_id -> asyncio.Event
cancel_events: Dict[int, asyncio.Event] = {}

# 活跃任务: task_id -> asyncio.Task
active_tasks: Dict[int, asyncio.Task] = {}
queue_dispatcher_task: Optional[asyncio.Task] = None
pause_requests: set[int] = set()
queue_dispatch_lock = asyncio.Lock()

# 进度写库节流：上次写库时间戳
_last_progress_db_write = 0.0


def _site_auth_enabled() -> bool:
    return bool(config.SITE_AUTH_PASSWORD)


def _sign_auth_payload(payload: str) -> str:
    return hmac.new(
        config.SITE_AUTH_SECRET.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()


def _make_auth_token(username: str) -> str:
    expires_at = int((datetime.utcnow() + timedelta(days=config.SITE_AUTH_SESSION_DAYS)).timestamp())
    payload = f"{username}:{expires_at}"
    encoded_payload = base64.urlsafe_b64encode(payload.encode("utf-8")).decode("ascii")
    return f"{encoded_payload}.{_sign_auth_payload(encoded_payload)}"


def _verify_auth_token(token: str) -> bool:
    try:
        encoded_payload, signature = token.split(".", 1)
        expected = _sign_auth_payload(encoded_payload)
        if not hmac.compare_digest(signature, expected):
            return False
        payload = base64.urlsafe_b64decode(encoded_payload.encode("ascii")).decode("utf-8")
        username, expires_at = payload.rsplit(":", 1)
        if username != config.SITE_AUTH_USERNAME:
            return False
        return int(expires_at) > int(datetime.utcnow().timestamp())
    except Exception:
        return False


def _is_secure_request(request: Request) -> bool:
    forwarded_proto = request.headers.get("x-forwarded-proto", "")
    return request.url.scheme == "https" or forwarded_proto.lower() == "https"


async def _find_live_task_for_novel(db: AsyncSession, novel_id: int) -> Optional[ExpandTask]:
    """Return an active task for this novel without breaking queued/paused work."""
    stmt = (
        select(ExpandTask)
        .where(ExpandTask.novel_id == novel_id)
        .where(ExpandTask.status.in_(["queued", "running", "pausing", "paused"]))
        .order_by(ExpandTask.created_at.desc())
    )
    result = await db.execute(stmt)
    tasks = result.scalars().all()

    live_task = None
    for task in tasks:
        if task.status in {"queued", "paused"}:
            live_task = task
            continue

        bg_task = active_tasks.get(task.id)
        if bg_task and not bg_task.done():
            live_task = task
            continue

        if task.status in {"running", "pausing"}:
            task.status = "interrupted"
            if not task.error_message:
                task.error_message = "Stale task auto-marked interrupted"
            task.updated_at = datetime.utcnow()

    if tasks:
        await db.commit()
    return live_task


async def _dispatch_next_task():
    async with queue_dispatch_lock:
        for task_id, task in list(active_tasks.items()):
            if task.done():
                active_tasks.pop(task_id, None)
                cancel_events.pop(task_id, None)

        if active_tasks:
            return None
        async with async_session() as db:
            running_result = await db.execute(
                select(ExpandTask.id).where(ExpandTask.status.in_(["running", "pausing"])).limit(1)
            )
            if running_result.scalar_one_or_none():
                return None

            stmt = (
                select(ExpandTask)
                .where(ExpandTask.status == "queued")
                .order_by(
                    ExpandTask.queue_priority.desc(),
                    func.coalesce(ExpandTask.queued_at, ExpandTask.created_at).asc(),
                    ExpandTask.id.asc(),
                )
                .limit(1)
            )
            result = await db.execute(stmt)
            task = result.scalar_one_or_none()
            if not task:
                return None

            task.status = "running"
            task.updated_at = datetime.utcnow()
            await db.commit()
            await db.refresh(task)

            cancel_event = asyncio.Event()
            cancel_events[task.id] = cancel_event
            resume_from_index = task.last_completed_index if task.last_completed_index is not None else -1
            bg_task = asyncio.create_task(
                expand_worker(task.id, cancel_event, resume_from_index=resume_from_index)
            )
            active_tasks[task.id] = bg_task
            logger.info(
                f"Dispatched queued task {task.id} for novel {task.novel_id} "
                f"(resume_from_index={resume_from_index})"
            )
            return task.id


async def _queue_dispatcher_loop():
    while True:
        try:
            await _dispatch_next_task()
        except Exception as e:
            logger.error(f"queue dispatcher loop error: {e}", exc_info=True)
        await asyncio.sleep(0.8)


def _is_fatal_api_error(error: Exception) -> bool:
    """这类错误通常是账号/反代/上游不可用，继续跑后续章节只会重复失败。"""
    text = str(error).lower()
    fatal_markers = (
        "chat upstream returned 403",
        "upstream_error",
        "invalid_api_key",
        "authentication_error",
        "permissiondenied",
        "permission denied",
        "model_not_found",
        "error code: 403",
        "http 403",
    )
    return any(marker in text for marker in fatal_markers)


# AI 在"无需扩写"时可能输出的通知短语（当系统角色配置为"还原器"时）
_NO_EXPANSION_NOTICE_SIGNALS = [
    "未检测到明确的省略",
    "未检测到需要扩写",
    "已按原文保留",
    "本章未检测到",
    "没有检测到删减",
    "无需扩写",
    "不需要扩写",
    "原文保留",
    "拒绝生成该内容",
    "拒绝处理此请求",
    "xAI核心安全准则",
    "我无法生成、扩展",
    "我无法生成、扩展、还原",
    "请提供明确以双方",
    "请提供**明确以成年",
    "此请求通过自定义",
    "属于典型的越狱",
]


def _is_no_expansion_notice(expanded: str, input_content: str) -> bool:
    """判断 AI 返回值是否属于"无需扩写"通知，而非真正的扩写结果。

    触发条件（任意满足）：
    1. 返回值与输入原文完全相同
    2. 返回值 <= 500字 且含已知通知短语
    3. 返回值不足原文的 30% 且含已知通知短语
    4. 返回值开头是"无需扩写"通知，后面又直接附带原文
    """
    if expanded == input_content:
        return True
    if not expanded:
        return True
    stripped = expanded.lstrip()
    head = stripped[:260]
    starts_with_notice = any(signal in head for signal in _NO_EXPANSION_NOTICE_SIGNALS)
    if starts_with_notice:
        if not input_content:
            return True
        anchors = []
        source = input_content.strip()
        if source:
            anchors.append(source[:120])
            if len(source) > 360:
                mid = len(source) // 2
                anchors.append(source[max(0, mid - 60):mid + 60])
                anchors.append(source[-120:])
        anchors = [anchor for anchor in anchors if len(anchor) >= 40]
        if not anchors or sum(1 for anchor in anchors if anchor in expanded) >= min(2, len(anchors)):
            return True
    if len(expanded) <= 500:
        for signal in _NO_EXPANSION_NOTICE_SIGNALS:
            if signal in expanded:
                return True
    if len(expanded) < len(input_content) * 0.3:
        for signal in _NO_EXPANSION_NOTICE_SIGNALS:
            if signal in expanded:
                return True
    return False


# ========== Pydantic 请求模型 ==========
class ExpandRequest(BaseModel):
    chapter_ids: Optional[List[int]] = None
    model: Optional[str] = None
    mode: str = "one_pass"  # legacy clients may send old values; server normalizes to one_pass
    quality: str = "balanced"  # legacy clients may send old values; server normalizes to balanced
    use_expanded_as_base: bool = False  # True = 基于已扩写内容再次扩写


class RewriteRequest(BaseModel):
    instruction: str = ""
    model: Optional[str] = None


class RewriteParagraphRequest(BaseModel):
    paragraph_index: int
    instruction: str = ""
    model: Optional[str] = None
    use_expanded: bool = True


class InsertPromptRequest(BaseModel):
    paragraph_index: int
    prompt: str
    model: Optional[str] = None


class SelectionExpandRequest(BaseModel):
    selected_text: str
    prompt: str = ""
    model: Optional[str] = None


class SaveContentRequest(BaseModel):
    content: str
    is_expanded: bool = True


class ModelTestRequest(BaseModel):
    model: Optional[str] = None
    prompt: str = "用一句话回答：你是谁？"


def _sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", name).strip() or "novel"


def _chapter_display_content(chapter: Chapter) -> str:
    return normalize_output_text(chapter.expanded_content if chapter.expanded_content else chapter.original_content)


def _chapter_weight(chapter: Chapter) -> int:
    content = _chapter_display_content(chapter)
    return max(len(content or ""), 1)


def _merge_context_summary(global_summary: str = "", prev_summary: str = "") -> str:
    parts = []
    if global_summary:
        parts.append(f"=== 全局人物/设定/伏笔摘要 ===\n{global_summary}")
    if prev_summary:
        parts.append(f"=== 前一章摘要 ===\n{prev_summary}")
    return "\n\n".join(parts)


async def _refresh_novel_global_summary(novel_id: int):
    """基于最近章节摘要生成滚动全局摘要，避免长篇上下文完全丢失。
    优化：按时间远近分配权重，近章给完整摘要，远章压缩。"""
    async with async_session() as db:
        stmt = (
            select(Novel)
            .options(selectinload(Novel.chapters))
            .where(Novel.id == novel_id)
        )
        result = await db.execute(stmt)
        novel = result.scalar_one_or_none()
        if not novel:
            return

        chapters_with_summary = sorted(
            (ch for ch in novel.chapters if ch.summary),
            key=lambda ch: ch.sort_order,
        )

        if not chapters_with_summary:
            novel.global_summary = None
        else:
            # 取最近16章（增加覆盖范围）
            recent = chapters_with_summary[-16:]
            summary_parts = []

            for i, chapter in enumerate(recent):
                # 越近的章节保留越多细节
                if i >= len(recent) - 3:
                    # 最近3章：完整摘要
                    summary_parts.append(f"{chapter.title}：{chapter.summary}")
                elif i >= len(recent) - 8:
                    # 中间5章：截取要点（取前300字）
                    truncated = chapter.summary[:300]
                    if len(chapter.summary) > 300:
                        truncated = truncated.rsplit('。', 1)[0] + '。' if '。' in truncated else truncated
                    summary_parts.append(f"{chapter.title}：{truncated}")
                else:
                    # 较早章节：只保留标题和首句
                    first_line = chapter.summary.split('\n')[0][:100]
                    summary_parts.append(f"{chapter.title}：{first_line}")

            merged = "\n".join(summary_parts)
            novel.global_summary = merged[-6000:]  # 增加到6000字符上限
        novel.updated_at = datetime.utcnow()
        await db.commit()


def _task_history_item(task: ExpandTask) -> Dict[str, object]:
    duration_seconds = None
    if task.created_at and task.updated_at:
        duration_seconds = max(int((task.updated_at - task.created_at).total_seconds()), 0)

    chapter_ids = None
    if task.chapter_ids_json:
        try:
            chapter_ids = json.loads(task.chapter_ids_json)
        except json.JSONDecodeError:
            chapter_ids = None

    return {
        "id": task.id,
        "novel_id": task.novel_id,
        "status": task.status,
        "model": task.model,
        "mode": task.mode,
        "quality": task.quality,
        "progress": task.progress,
        "total_chapters": task.total_chapters,
        "completed_chapters": task.completed_chapters,
        "failed_chapters": task.failed_chapters,
        "skipped_chapters": task.skipped_chapters,
        "queue_priority": task.queue_priority or 0,
        "queued_at": task.queued_at.isoformat() if task.queued_at else None,
        "last_completed_index": task.last_completed_index,
        "current_chapter_title": task.current_chapter_title,
        "chapter_ids": chapter_ids,
        "use_expanded_as_base": task.use_expanded_as_base,
        "error_message": task.error_message,
        "duration_seconds": duration_seconds,
        "created_at": task.created_at.isoformat() if task.created_at else None,
        "updated_at": task.updated_at.isoformat() if task.updated_at else None,
    }


def _build_export_text(novel: Novel, separator_style: str = "classic") -> str:
    if separator_style == "minimal":
        sep = "\n"
    elif separator_style == "double":
        sep = "\n" + ("=" * 60) + "\n"
    else:
        sep = "\n" + ("-" * 40) + "\n"

    parts = [novel.title, "=" * 40, ""]
    for ch in novel.chapters:
        parts.append(ch.title)
        parts.append("")
        parts.append(_chapter_display_content(ch))
        parts.append(sep)
    return "\n".join(parts).strip() + "\n"


def _build_docx(path: str, novel: Novel):
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    doc.add_heading(novel.title, level=0)
    for ch in novel.chapters:
        doc.add_heading(ch.title, level=1)
        for para in split_into_paragraphs(_chapter_display_content(ch)):
            doc.add_paragraph(para)
    doc.save(path)


def _build_epub(path: str, novel: Novel):
    if epub is None:
        raise RuntimeError("ebooklib is not installed")
    book = epub.EpubBook()
    book.set_identifier(f"novel-expander-{novel.id}")
    book.set_title(novel.title)
    book.set_language("zh-CN")
    book.add_author("Novel Expander")

    spine = ["nav"]
    toc = []
    for idx, ch in enumerate(novel.chapters, start=1):
        item = epub.EpubHtml(
            title=ch.title,
            file_name=f"chapter_{idx}.xhtml",
            lang="zh-CN",
        )
        paragraphs = "".join(
            f"<p>{escape(para)}</p>" for para in split_into_paragraphs(_chapter_display_content(ch))
        )
        item.content = f"<h1>{escape(ch.title)}</h1>{paragraphs}"
        book.add_item(item)
        toc.append(item)
        spine.append(item)

    book.toc = tuple(toc)
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(path, book)


# ========== SSE 广播 ==========
async def broadcast_sse(novel_id: int, event: str, data: dict):
    """向所有订阅该小说的客户端广播 SSE 事件"""
    if novel_id in sse_subscribers:
        msg = {"event": event, "data": json.dumps(data, ensure_ascii=False)}
        dead_queues = []
        for queue in sse_subscribers[novel_id]:
            try:
                await queue.put(msg)
            except Exception:
                dead_queues.append(queue)
        for q in dead_queues:
            try:
                sse_subscribers[novel_id].remove(q)
            except ValueError:
                pass


# ========== Token 池检查辅助函数 ==========
async def _check_token_pool_available() -> bool:
    """检查 token 池是否可用"""
    admin_url = config.API_BASE.rstrip('/').replace('/v1', '/v1/admin/tokens')
    try:
        async with httpx.AsyncClient(timeout=5.0) as http_client:
            resp = await http_client.get(
                admin_url,
                headers={"Authorization": f"Bearer {config.ADMIN_API_KEY}"},
            )
            if resp.status_code != 200:
                return True  # 查询失败不阻止
            data = resp.json()
            tokens_dict = data.get("tokens", {})
            for pool_tokens in tokens_dict.values():
                for tk in pool_tokens:
                    if tk.get("status") == "active":
                        return True
            return False
    except Exception:
        return True  # 查询失败不阻止


async def _wait_for_token_pool() -> bool:
    """等待 token 池恢复可用，返回是否恢复"""
    waited = 0.0
    while waited < config.TOKEN_POOL_WAIT_MAX:
        await asyncio.sleep(config.TOKEN_POOL_CHECK_INTERVAL)
        waited += config.TOKEN_POOL_CHECK_INTERVAL
        if await _check_token_pool_available():
            logger.info(f"Token pool recovered after {waited:.0f}s")
            return True
    logger.error(f"Token pool still unavailable after {waited:.0f}s")
    return False


async def _chapter_cooldown(
    novel_id: int,
    task_id: int,
    chapter_id: int,
    chapter_title: str,
    completed_count: int,
    failed_count: int,
    skipped_count: int,
    total_chapters: int,
    completed_weight: int,
    total_weight: int,
    cancel_event: Optional[asyncio.Event] = None,
):
    """章节之间额外冷却，避免批量任务连续打请求。"""
    delay = config.INTER_CHAPTER_DELAY_SECONDS
    if delay <= 0:
        return

    await broadcast_sse(novel_id, "progress", {
        "task_id": task_id,
        "chapter_id": chapter_id,
        "chapter_title": chapter_title,
        "status": "waiting",
        "message": f"⏳ 章节间冷却中，等待 {int(delay)} 秒后继续...",
        "chapter_progress": 0.0,
        "overall_progress": completed_weight / total_weight if total_weight else 0.0,
        "completed_chapters": completed_count,
        "failed_chapters": failed_count,
        "skipped_chapters": skipped_count,
        "total_chapters": total_chapters,
    })
    if cancel_event is None:
        await asyncio.sleep(delay)
        return

    # Allow responsive cancel during cooldown.
    waited = 0.0
    step = min(1.0, delay)
    while waited < delay:
        if cancel_event.is_set():
            return
        await asyncio.sleep(step)
        waited += step




# ========== 进度更新节流 ==========
async def _debounced_progress_update(
    task_id: int,
    chapter_id: int,
    ch_progress: float,
    overall: float,
    message: str,
    novel_id: int,
    chapter_title: str,
    completed: int,
    failed: int,
    skipped: int,
    total: int,
    current_weight: int = 1,
):
    """节流的进度数据库更新，避免频繁写库"""
    global _last_progress_db_write
    now = time.time()

    # 始终广播 SSE（前端需要实时更新）
    await broadcast_sse(novel_id, "progress", {
        "task_id": task_id,
        "chapter_id": chapter_id,
        "chapter_title": chapter_title,
        "status": "expanding",
        "chapter_progress": ch_progress,
        "overall_progress": overall,
        "completed_chapters": completed,
        "failed_chapters": failed,
        "skipped_chapters": skipped,
        "total_chapters": total,
        "message": message,
    })

    # 节流写库（每 N 秒最多一次）
    if now - _last_progress_db_write >= config.PROGRESS_DEBOUNCE_SECONDS:
        _last_progress_db_write = now
        async with async_session() as prog_db:
            pg_ch = await prog_db.get(Chapter, chapter_id)
            if pg_ch:
                pg_ch.progress = ch_progress
                pg_ch.updated_at = datetime.utcnow()
            pg_task = await prog_db.get(ExpandTask, task_id)
            if pg_task:
                pg_task.progress = overall
                pg_task.current_chapter_title = chapter_title
                pg_task.updated_at = datetime.utcnow()
            await prog_db.commit()


# ========== Lifespan（替代废弃的 on_event） ==========
@asynccontextmanager
async def lifespan(app):
    # === startup ===
    logger.info("Initializing database...")
    await init_db()
    await migrate_db()  # 数据库迁移：为已有表添加新增列
    settings_manager.init_settings()  # 加载运行时设置
    global queue_dispatcher_task

    # 将中断的任务标记为 interrupted（不是 failed，方便恢复）
    async with async_session() as db:
        stmt = (
            update(ExpandTask)
            .where(ExpandTask.status.in_(["running", "pausing"]))
            .values(status="interrupted", error_message="Server restarted")
        )
        await db.execute(stmt)

        # 同时重置运行中的章节状态
        stmt2 = (
            update(Chapter)
            .where(Chapter.status.in_(["analyzing", "expanding"]))
            .values(status="pending", progress=0.0)
        )
        await db.execute(stmt2)
        await db.commit()

    queue_dispatcher_task = asyncio.create_task(_queue_dispatcher_loop())

    # 清理旧导出文件
    export_dir = os.path.join(config.DATA_DIR, "exports")
    if os.path.isdir(export_dir):
        for f in os.listdir(export_dir):
            try:
                os.remove(os.path.join(export_dir, f))
            except Exception:
                pass

    logger.info(f"Server ready at http://{config.HOST}:{config.PORT}")

    yield  # 运行中

    # === shutdown — 优雅停机 ===
    logger.info("Shutting down, cancelling active tasks...")
    if queue_dispatcher_task:
        queue_dispatcher_task.cancel()
    for task_id, cancel_evt in list(cancel_events.items()):
        cancel_evt.set()
    # 等待活跃任务结束（最多10秒）
    for task_id, task in list(active_tasks.items()):
        try:
            await asyncio.wait_for(asyncio.shield(task), timeout=10.0)
        except (asyncio.TimeoutError, asyncio.CancelledError, Exception):
            pass
    logger.info("Shutdown complete")


# ========== FastAPI 应用 ==========
app = FastAPI(title="Novel Expander API", version="2.0.0", lifespan=lifespan)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def site_auth(request: Request, call_next):
    if not _site_auth_enabled():
        return await call_next(request)

    path = request.url.path
    public_paths = {"/login", "/api/login", "/api/logout"}
    if path in public_paths or path.startswith("/static/"):
        return await call_next(request)

    token = request.cookies.get(config.SITE_AUTH_COOKIE, "")
    if _verify_auth_token(token):
        return await call_next(request)

    if path.startswith("/api/"):
        return JSONResponse({"detail": "Not authenticated"}, status_code=401)

    next_path = path
    if request.url.query:
        next_path = f"{next_path}?{request.url.query}"
    return RedirectResponse(url=f"/login?next={next_path}", status_code=303)


@app.middleware("http")
async def log_requests(request: Request, call_next):
    response = await call_next(request)
    if response.status_code >= 400:
        logger.warning(f"HTTP {response.status_code}: {request.method} {request.url.path}")
    return response


# ========== 静态文件 ==========
static_dir = os.path.join(config.BASE_DIR, "static")
if os.path.isdir(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")


@app.get("/login")
async def login_page():
    login_path = os.path.join(static_dir, "login.html")
    if os.path.exists(login_path):
        return FileResponse(login_path)
    return {"message": "Login page not found."}


@app.post("/api/login")
async def login(
    request: Request,
    username: str = Form(""),
    password: str = Form(""),
    next: str = Form("/"),
):
    if not _site_auth_enabled():
        return RedirectResponse(url="/", status_code=303)

    valid_username = hmac.compare_digest(username, config.SITE_AUTH_USERNAME)
    valid_password = hmac.compare_digest(password, config.SITE_AUTH_PASSWORD)
    if not (valid_username and valid_password):
        return RedirectResponse(url="/login?error=1", status_code=303)

    safe_next = next if next.startswith("/") and not next.startswith("//") else "/"
    response = RedirectResponse(url=safe_next, status_code=303)
    response.set_cookie(
        key=config.SITE_AUTH_COOKIE,
        value=_make_auth_token(username),
        max_age=config.SITE_AUTH_SESSION_DAYS * 24 * 60 * 60,
        httponly=True,
        secure=_is_secure_request(request),
        samesite="lax",
        path="/",
    )
    return response


@app.post("/api/logout")
async def logout():
    response = RedirectResponse(url="/login", status_code=303)
    response.delete_cookie(config.SITE_AUTH_COOKIE, path="/")
    return response


@app.get("/")
async def index():
    index_path = os.path.join(static_dir, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Novel Expander API is running. Static files not found."}


@app.get("/api/health")
async def health_check():
    return {
        "status": "ok",
        "time": datetime.utcnow().isoformat(),
        "active_tasks": len(active_tasks),
    }


@app.post("/api/model-test")
async def model_test(req: ModelTestRequest):
    """Smoke test: verifies the upstream /v1/chat/completions really works."""
    t0 = time.time()
    text = await chat_completion(
        messages=[
            {"role": "system", "content": "你是一个简洁可靠的助手。"},
            {"role": "user", "content": req.prompt},
        ],
        model=req.model,
        stream=False,
        temperature=0.2,
    )
    latency_ms = int((time.time() - t0) * 1000)
    return {
        "ok": True,
        "api_base": config.API_BASE,
        "requested_model": req.model,
        "effective_model_candidates": config.get_model_candidates(req.model),
        "latency_ms": latency_ms,
        "text": (text or "").strip(),
    }


# ========== 设置管理路由 ==========

@app.get("/api/settings")
async def get_settings():
    """获取所有运行时设置"""
    return {
        "settings": settings_manager.get_all(),
        "meta": settings_manager.get_meta(),
    }


@app.put("/api/settings")
async def update_settings(request: Request):
    """批量更新运行时设置"""
    body = await request.json()
    changes = body.get("settings", body)
    if not isinstance(changes, dict):
        raise HTTPException(400, "Invalid settings payload")
    updated = settings_manager.update(changes)
    return {
        "updated": updated,
        "settings": settings_manager.get_all(),
    }


@app.post("/api/settings/reset")
async def reset_settings():
    """重置所有设置为默认值"""
    settings_manager.reset_to_defaults()
    return {
        "settings": settings_manager.get_all(),
        "meta": settings_manager.get_meta(),
    }


# ========== 提示词设置路由 ==========

@app.get("/api/prompts")
async def get_prompts():
    """获取所有可编辑提示词模板"""
    return get_prompt_settings()


@app.put("/api/prompts")
async def update_prompts(request: Request):
    """批量更新提示词模板"""
    body = await request.json()
    changes = body.get("prompts", body)
    if not isinstance(changes, dict):
        raise HTTPException(400, "Invalid prompts payload")
    updated = update_prompt_settings(changes)
    return {
        "updated": updated,
        **get_prompt_settings(),
    }


@app.post("/api/prompts/reset")
async def reset_prompts(request: Request):
    """恢复全部或部分提示词为代码默认值"""
    try:
        body = await request.json()
    except Exception:
        body = {}
    keys = body.get("keys") if isinstance(body, dict) else None
    if keys is not None and not isinstance(keys, list):
        raise HTTPException(400, "Invalid prompt keys")
    reset_prompt_settings(keys)
    return get_prompt_settings()


# ========== API Profiles 路由 ==========

@app.get("/api/profiles")
async def get_profiles():
    """获取所有 API 配置"""
    return {
        "profiles": settings_manager.get_profiles(),
        "active_profile_id": settings_manager.get_active_profile_id(),
    }


@app.post("/api/profiles")
async def create_profile(request: Request):
    """新增 API 配置"""
    body = await request.json()
    profile = settings_manager.add_profile(body)
    # Never return secrets to clients.
    redacted = next((p for p in settings_manager.get_profiles() if p.get("id") == profile.get("id")), None)
    return {"profile": redacted or {"id": profile.get("id")}}


@app.put("/api/profiles/{profile_id}")
async def update_profile(profile_id: str, request: Request):
    """更新指定 API 配置"""
    body = await request.json()
    profile = settings_manager.update_profile(profile_id, body)
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    redacted = next((p for p in settings_manager.get_profiles() if p.get("id") == profile_id), None)
    return {"profile": redacted or {"id": profile_id}}


@app.delete("/api/profiles/{profile_id}")
async def delete_profile(profile_id: str):
    """删除指定 API 配置"""
    ok = settings_manager.delete_profile(profile_id)
    if not ok:
        raise HTTPException(status_code=400, detail="Cannot delete (at least one profile required)")
    return {
        "message": "Profile deleted",
        "active_profile_id": settings_manager.get_active_profile_id(),
    }


@app.post("/api/profiles/{profile_id}/switch")
async def switch_profile(profile_id: str):
    """切换激活的 API 配置"""
    profile = settings_manager.switch_profile(profile_id)
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    redacted = next((p for p in settings_manager.get_profiles() if p.get("id") == profile_id), None)
    return {"profile": redacted or {"id": profile_id}}


# ========== Token 状态路由 ==========

@app.get("/api/token-status")
async def get_token_status():
    """获取 API token 池状态"""
    admin_url = config.API_BASE.rstrip('/').replace('/v1', '/v1/admin/tokens')
    try:
        async with httpx.AsyncClient(timeout=10.0) as http_client:
            resp = await http_client.get(
                admin_url,
                headers={
                    "Authorization": f"Bearer {config.ADMIN_API_KEY}",
                    "Accept": "application/json",
                },
            )
            resp.raise_for_status()
            data = resp.json()
    except Exception as e:
        logger.warning(f"Failed to fetch token status: {e}")
        return {"error": str(e), "available": True}  # 查询失败时不阻止扩写

    tokens_dict = data.get("tokens", {})
    consumed_mode = data.get("consumed_mode_enabled", False)

    total = 0
    active = 0
    cooling = 0
    failed = 0
    chat_remaining = 0

    for pool_name, pool_tokens in tokens_dict.items():
        for tk in pool_tokens:
            total += 1
            status = tk.get("status", "")
            if status == "active":
                active += 1
                chat_remaining += tk.get("quota", 0)
            elif status == "cooling":
                cooling += 1
            else:
                failed += 1

    return {
        "total": total,
        "active": active,
        "cooling": cooling,
        "failed": failed,
        "chat_remaining": chat_remaining,
        "consumed_mode": consumed_mode,
        "available": active > 0,
    }


# ========== 小说管理路由 ==========

@app.post("/api/novels/upload")
async def upload_novel(file: UploadFile = File(...), db: AsyncSession = Depends(get_db)):
    """上传小说文件"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # 检查扩展名
    if not file.filename.lower().endswith(".txt"):
        raise HTTPException(status_code=400, detail="Only .txt files are supported")

    try:
        raw_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {e}")

    if not raw_bytes:
        raise HTTPException(status_code=400, detail="File is empty")

    # 限制文件大小（50MB）
    if len(raw_bytes) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 50MB)")

    # 解析小说
    try:
        parsed = parse_novel_from_bytes(raw_bytes, file.filename)
    except Exception as e:
        logger.error(f"Failed to parse novel: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to parse novel: {e}")

    if not parsed.get("chapters"):
        raise HTTPException(status_code=400, detail="No chapters found in the file")

    # 存入数据库
    novel = Novel(
        title=parsed["title"],
        original_filename=file.filename,
    )
    db.add(novel)
    await db.flush()  # 获取 novel.id

    for i, ch in enumerate(parsed["chapters"]):
        chapter = Chapter(
            novel_id=novel.id,
            title=ch["title"],
            sort_order=i,
            original_content=ch["content"],
            status="pending",
        )
        db.add(chapter)

    await db.commit()
    await db.refresh(novel)

    logger.info(f"Uploaded novel: {novel.title} with {len(parsed['chapters'])} chapters")

    return {
        "id": novel.id,
        "title": novel.title,
        "chapter_count": len(parsed["chapters"]),
    }


@app.get("/api/novels")
async def list_novels(db: AsyncSession = Depends(get_db)):
    """列出所有小说"""
    stmt = select(Novel).order_by(Novel.created_at.desc())
    result = await db.execute(stmt)
    novels = result.scalars().all()

    novels_data = []
    for novel in novels:
        # 获取章节统计
        ch_stmt = select(
            func.count(Chapter.id),
            func.count(Chapter.expanded_content),
        ).where(Chapter.novel_id == novel.id)
        ch_result = await db.execute(ch_stmt)
        row = ch_result.one()
        total_chapters = row[0]
        # expanded_content 不为 null 的数量
        expanded_stmt = (
            select(func.count(Chapter.id))
            .where(Chapter.novel_id == novel.id)
            .where(Chapter.expanded_content.isnot(None))
        )
        expanded_result = await db.execute(expanded_stmt)
        expanded_count = expanded_result.scalar()

        novels_data.append({
            "id": novel.id,
            "title": novel.title,
            "chapter_count": total_chapters,
            "created_at": novel.created_at.isoformat() if novel.created_at else None,
            "has_expanded": expanded_count > 0,
        })

    return {"novels": novels_data}


@app.get("/api/novels/{novel_id}")
async def get_novel(novel_id: int, db: AsyncSession = Depends(get_db)):
    """获取小说详情"""
    stmt = select(Novel).options(selectinload(Novel.chapters)).where(Novel.id == novel_id)
    result = await db.execute(stmt)
    novel = result.scalar_one_or_none()

    if not novel:
        raise HTTPException(status_code=404, detail="Novel not found")

    chapters_data = []
    for ch in novel.chapters:
        chapters_data.append({
            "id": ch.id,
            "title": ch.title,
            "sort_order": ch.sort_order,
            "status": ch.status,
            "progress": ch.progress,
            "has_expanded": ch.expanded_content is not None,
            "original_word_count": len(ch.original_content) if ch.original_content else 0,
            "expanded_word_count": len(ch.expanded_content) if ch.expanded_content else 0,
            "error_message": ch.error_message,  # 新增
            "skipped": ch.skipped,  # 新增
        })

    return {
        "id": novel.id,
        "title": novel.title,
        "global_summary": novel.global_summary,
        "chapters": chapters_data,
    }


@app.delete("/api/novels/{novel_id}")
async def delete_novel(novel_id: int, db: AsyncSession = Depends(get_db)):
    """删除小说"""
    stmt = select(Novel).where(Novel.id == novel_id)
    result = await db.execute(stmt)
    novel = result.scalar_one_or_none()

    if not novel:
        raise HTTPException(status_code=404, detail="Novel not found")

    await db.delete(novel)
    await db.commit()

    return {"message": "Novel deleted"}


# ========== 章节路由 ==========

@app.get("/api/novels/{novel_id}/chapters/{chapter_id}")
async def get_chapter(novel_id: int, chapter_id: int, db: AsyncSession = Depends(get_db)):
    """获取章节内容"""
    stmt = select(Chapter).where(Chapter.id == chapter_id, Chapter.novel_id == novel_id)
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()

    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")

    # 分段
    paragraphs = []
    original_content = normalize_output_text(chapter.original_content)
    expanded_content = normalize_output_text(chapter.expanded_content)
    expanded_content_prev = normalize_output_text(chapter.expanded_content_prev)

    if original_content:
        paras = split_into_paragraphs(original_content)
        paragraphs = [{"index": i, "text": p} for i, p in enumerate(paras)]

    expanded_paragraphs = None
    if expanded_content:
        exp_paras = split_into_paragraphs(expanded_content)
        expanded_paragraphs = [{"index": i, "text": p} for i, p in enumerate(exp_paras)]

    return {
        "id": chapter.id,
        "title": chapter.title,
        "original_content": original_content,
        "expanded_content": expanded_content,
        "expanded_content_prev": expanded_content_prev,  # 新增
        "status": chapter.status,
        "skipped": chapter.skipped,  # 新增
        "error_message": chapter.error_message,  # 新增
        "summary": chapter.summary,  # 新增
        "paragraphs": paragraphs,
        "expanded_paragraphs": expanded_paragraphs,
    }


# ========== 扩写路由 ==========

@app.post("/api/novels/{novel_id}/expand")
async def start_expand(novel_id: int, body: ExpandRequest, db: AsyncSession = Depends(get_db)):
    """开始扩写任务"""
    # 验证小说存在
    stmt = select(Novel).where(Novel.id == novel_id)
    result = await db.execute(stmt)
    novel = result.scalar_one_or_none()
    if not novel:
        raise HTTPException(status_code=404, detail="Novel not found")

    # 检查是否已有运行中的任务
    live_task = await _find_live_task_for_novel(db, novel_id)
    if live_task:
        raise HTTPException(status_code=409, detail="An expand task is already queued, pausing, paused, or running for this novel")

    # 确定要处理的章节
    chapter_ids_json = None
    if body.chapter_ids:
        chapter_ids_json = json.dumps(body.chapter_ids)

    # 计算总章节数
    if body.chapter_ids:
        total = len(body.chapter_ids)
    else:
        count_stmt = select(func.count(Chapter.id)).where(Chapter.novel_id == novel_id)
        count_result = await db.execute(count_stmt)
        total = count_result.scalar()

    # 创建任务
    task = ExpandTask(
        novel_id=novel_id,
        status="queued",
        model=body.model or config.DEFAULT_MODEL,
        mode="one_pass",
        quality="balanced",
        total_chapters=total,
        chapter_ids_json=chapter_ids_json,
        use_expanded_as_base=body.use_expanded_as_base,
        queued_at=datetime.utcnow(),
    )
    db.add(task)
    await db.commit()
    await db.refresh(task)

    logger.info(f"Queued expand task {task.id} for novel {novel_id}")
    await _dispatch_next_task()
    await db.refresh(task)
    return {"task_id": task.id, "status": task.status}


@app.get("/api/novels/{novel_id}/expand/stream")
async def expand_stream(novel_id: int):
    """SSE 流推送扩写进度"""
    queue = asyncio.Queue()
    if novel_id not in sse_subscribers:
        sse_subscribers[novel_id] = []
    sse_subscribers[novel_id].append(queue)

    async def event_generator():
        try:
            # If a task finished before the client connected (very fast skip/completion),
            # push a recent snapshot so the UI can stop "spinning".
            try:
                async with async_session() as snap_db:
                    stmt = (
                        select(ExpandTask)
                        .where(ExpandTask.novel_id == novel_id)
                        .order_by(ExpandTask.created_at.desc())
                        .limit(1)
                    )
                    result = await snap_db.execute(stmt)
                    last_task = result.scalars().first()
                    if last_task and last_task.created_at:
                        age = (datetime.utcnow() - last_task.created_at).total_seconds()
                        if age < 120 and last_task.status not in {"queued", "running", "pausing"}:
                            yield {
                                "event": "task_done",
                                "data": json.dumps(_task_history_item(last_task), ensure_ascii=False),
                            }
            except Exception:
                pass

            while True:
                try:
                    msg = await asyncio.wait_for(queue.get(), timeout=10.0)
                    yield msg
                except asyncio.TimeoutError:
                    yield {"event": "heartbeat", "data": "{}"}
        except asyncio.CancelledError:
            pass
        finally:
            if novel_id in sse_subscribers:
                try:
                    sse_subscribers[novel_id].remove(queue)
                except ValueError:
                    pass
                if not sse_subscribers[novel_id]:
                    del sse_subscribers[novel_id]

    return EventSourceResponse(event_generator())


@app.post("/api/novels/{novel_id}/expand/cancel")
async def cancel_expand(novel_id: int, db: AsyncSession = Depends(get_db)):
    """取消扩写任务"""
    # 查找运行中的任务
    task = await _find_live_task_for_novel(db, novel_id)

    if not task:
        # Idempotent cancel: if no live task, return the latest task snapshot (if any)
        # so the frontend can stop retrying "cancel" forever.
        stmt = (
            select(ExpandTask)
            .where(ExpandTask.novel_id == novel_id)
            .order_by(ExpandTask.created_at.desc())
            .limit(1)
        )
        result = await db.execute(stmt)
        last_task = result.scalars().first()
        if last_task:
            await broadcast_sse(novel_id, "task_done", _task_history_item(last_task))
            return {"message": "No running task; returned latest task snapshot", "task": _task_history_item(last_task)}
        return {"message": "No running task"}

    # 发送取消信号
    if task.id in cancel_events:
        cancel_events[task.id].set()
    # Best-effort immediate stop (e.g., during long upstream call/sleep).
    bg = active_tasks.get(task.id)
    if bg and not bg.done():
        bg.cancel()

    task.status = "cancelled"
    task.updated_at = datetime.utcnow()
    await db.commit()

    # 广播取消事件
    await broadcast_sse(novel_id, "task_done", {
        "task_id": task.id,
        "status": "cancelled",
    })

    logger.info(f"Cancelled task {task.id}")

    return {"message": "Task cancelled"}


# ========== 恢复/重试/撤销 路由 ==========

@app.post("/api/novels/{novel_id}/expand/resume")
async def resume_expand(novel_id: int, db: AsyncSession = Depends(get_db)):
    """恢复中断的扩写任务"""
    # 查找 interrupted 状态的任务
    stmt = select(ExpandTask).where(
        ExpandTask.novel_id == novel_id,
        ExpandTask.status == "interrupted"
    ).order_by(ExpandTask.created_at.desc()).limit(1)
    result = await db.execute(stmt)
    task = result.scalars().first()

    if not task:
        raise HTTPException(status_code=404, detail="No interrupted task found")

    # 重置状态并排队
    task.status = "queued"
    task.queue_priority = 0
    task.queued_at = datetime.utcnow()
    task.error_message = None
    await db.commit()
    await db.refresh(task)

    logger.info(f"Queued resumed task {task.id} from index {task.last_completed_index}")
    await _dispatch_next_task()
    await db.refresh(task)
    return {"task_id": task.id, "resumed_from_index": task.last_completed_index, "status": task.status}


@app.post("/api/novels/{novel_id}/expand/retry-failed")
async def retry_failed_chapters(novel_id: int, db: AsyncSession = Depends(get_db)):
    """重试最近一次任务中失败的章节"""
    # 找最近完成/失败的任务
    stmt = select(ExpandTask).where(
        ExpandTask.novel_id == novel_id,
        ExpandTask.status.in_(["completed", "failed"]),
        ExpandTask.failed_chapter_ids_json.isnot(None),
    ).order_by(ExpandTask.created_at.desc())
    result = await db.execute(stmt)
    old_task = result.scalar_one_or_none()

    if not old_task or not old_task.failed_chapter_ids_json:
        raise HTTPException(status_code=404, detail="No failed chapters found")

    failed_ids = json.loads(old_task.failed_chapter_ids_json)
    if not failed_ids:
        raise HTTPException(status_code=404, detail="No failed chapters to retry")

    # 创建新任务只处理失败章节
    new_task = ExpandTask(
        novel_id=novel_id,
        status="queued",
        model=old_task.model,
        mode="one_pass",
        quality="balanced",
        total_chapters=len(failed_ids),
        chapter_ids_json=json.dumps(failed_ids),
        use_expanded_as_base=old_task.use_expanded_as_base,
        queued_at=datetime.utcnow(),
    )
    db.add(new_task)
    await db.commit()
    await db.refresh(new_task)

    logger.info(f"Queued retry {len(failed_ids)} failed chapters for novel {novel_id}, new task {new_task.id}")
    await _dispatch_next_task()
    await db.refresh(new_task)
    return {"task_id": new_task.id, "retrying_chapters": len(failed_ids), "status": new_task.status}


@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/undo")
async def undo_expansion(novel_id: int, chapter_id: int, db: AsyncSession = Depends(get_db)):
    """撤销最近一次扩写，恢复上一版本"""
    stmt = select(Chapter).where(Chapter.id == chapter_id, Chapter.novel_id == novel_id)
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    if not chapter.expanded_content_prev:
        raise HTTPException(status_code=400, detail="No previous version to restore")

    # 交换 current 和 prev
    chapter.expanded_content, chapter.expanded_content_prev = (
        chapter.expanded_content_prev, chapter.expanded_content
    )
    chapter.updated_at = datetime.utcnow()
    await db.commit()

    return {"message": "Undo successful", "chapter_id": chapter_id}


@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/save-content")
async def save_chapter_content(
    novel_id: int,
    chapter_id: int,
    body: SaveContentRequest,
    db: AsyncSession = Depends(get_db),
):
    """手动保存章节内容，支持撤销到上一版本"""
    stmt = select(Chapter).where(Chapter.id == chapter_id, Chapter.novel_id == novel_id)
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")

    content = normalize_output_text((body.content or "").strip())
    if not content:
        raise HTTPException(status_code=400, detail="Content cannot be empty")

    if body.is_expanded:
        if chapter.expanded_content != content:
            chapter.expanded_content_prev = chapter.expanded_content
            chapter.expanded_content = content
        chapter.status = "completed"
        chapter.skipped = False
    else:
        chapter.original_content = content
        if chapter.status == "failed":
            chapter.status = "pending"

    summary_source = content if body.is_expanded else (chapter.expanded_content or content)
    chapter.summary = build_local_chapter_summary(chapter.title, summary_source)
    chapter.error_message = None
    chapter.progress = 1.0
    chapter.updated_at = datetime.utcnow()
    await db.commit()
    await _refresh_novel_global_summary(novel_id)

    return {"message": "Content saved", "chapter_id": chapter_id, "is_expanded": body.is_expanded}


# ========== 预估/状态查询 路由 ==========

@app.get("/api/novels/{novel_id}/expand/estimate")
async def estimate_expand(novel_id: int, chapter_ids: str = None, db: AsyncSession = Depends(get_db)):
    """预估扩写时间和 token 消耗"""
    if chapter_ids:
        ids = json.loads(chapter_ids)
        stmt = select(Chapter).where(Chapter.novel_id == novel_id, Chapter.id.in_(ids))
    else:
        stmt = select(Chapter).where(Chapter.novel_id == novel_id)
    result = await db.execute(stmt)
    chapters = result.scalars().all()

    total_chars = sum(len(ch.original_content) for ch in chapters if ch.original_content)
    avg_time_per_1k_chars = 15  # 估算：每1000字约15秒
    estimated_seconds = (total_chars / 1000) * avg_time_per_1k_chars
    estimated_tokens = int(total_chars * config.TOKEN_PER_CHAR_ZH * 3)  # 输入+输出约3倍

    return {
        "chapter_count": len(chapters),
        "total_chars": total_chars,
        "estimated_seconds": int(estimated_seconds),
        "estimated_minutes": round(estimated_seconds / 60, 1),
        "estimated_tokens": estimated_tokens,
    }


@app.get("/api/novels/{novel_id}/expand/interrupted")
async def check_interrupted(novel_id: int, db: AsyncSession = Depends(get_db)):
    """检查是否有中断的任务可以恢复"""
    stmt = select(ExpandTask).where(
        ExpandTask.novel_id == novel_id,
        ExpandTask.status == "interrupted"
    ).order_by(ExpandTask.created_at.desc()).limit(1)
    result = await db.execute(stmt)
    task = result.scalars().first()

    if not task:
        return {"has_interrupted": False}

    return {
        "has_interrupted": True,
        "task_id": task.id,
        "last_completed_index": task.last_completed_index,
        "total_chapters": task.total_chapters,
        "completed_chapters": task.completed_chapters,
        "failed_chapters": task.failed_chapters,
        "skipped_chapters": task.skipped_chapters,
        "model": task.model,
        "mode": task.mode,
        "created_at": task.created_at.isoformat() if task.created_at else None,
    }


@app.get("/api/novels/{novel_id}/tasks")
async def get_task_history(novel_id: int, limit: int = 20, db: AsyncSession = Depends(get_db)):
    limit = max(1, min(limit, 100))
    status_rank = case(
        (ExpandTask.status.in_(["running", "pausing"]), 0),
        (ExpandTask.status == "queued", 1),
        (ExpandTask.status == "paused", 2),
        else_=3,
    )
    stmt = (
        select(ExpandTask)
        .where(ExpandTask.novel_id == novel_id)
        .order_by(
            status_rank,
            ExpandTask.queue_priority.desc(),
            ExpandTask.updated_at.desc(),
            ExpandTask.created_at.desc(),
            ExpandTask.id.desc(),
        )
        .limit(limit)
    )
    result = await db.execute(stmt)
    tasks = result.scalars().all()
    return {"tasks": [_task_history_item(task) for task in tasks]}


@app.get("/api/tasks/queue")
async def list_global_queue(limit: int = 80, db: AsyncSession = Depends(get_db)):
    limit = max(1, min(limit, 200))
    status_rank = case(
        (ExpandTask.status.in_(["running", "pausing"]), 0),
        (ExpandTask.status == "queued", 1),
        (ExpandTask.status == "paused", 2),
        else_=3,
    )
    stmt = (
        select(ExpandTask)
        .order_by(
            status_rank,
            ExpandTask.queue_priority.desc(),
            ExpandTask.updated_at.desc(),
            ExpandTask.created_at.desc(),
            ExpandTask.id.desc(),
        )
        .limit(limit)
    )
    result = await db.execute(stmt)
    tasks = result.scalars().all()
    novel_ids = {task.novel_id for task in tasks}
    titles = {}
    if novel_ids:
        novels_result = await db.execute(select(Novel.id, Novel.title).where(Novel.id.in_(novel_ids)))
        titles = {novel_id: title for novel_id, title in novels_result.all()}
    items = []
    for task in tasks:
        item = _task_history_item(task)
        item["novel_title"] = titles.get(task.novel_id, f"Novel {task.novel_id}")
        items.append(item)
    return {"tasks": items}


@app.delete("/api/tasks/history")
async def clear_task_history(db: AsyncSession = Depends(get_db)):
    """清空历史任务，保留正在运行、排队和暂停相关任务。"""
    protected_statuses = ["queued", "running", "pausing", "paused"]
    protected_ids = set(active_tasks.keys())
    stmt = delete(ExpandTask).where(~ExpandTask.status.in_(protected_statuses))
    if protected_ids:
        stmt = stmt.where(~ExpandTask.id.in_(protected_ids))
    result = await db.execute(stmt)
    await db.commit()
    return {"deleted": result.rowcount or 0}


@app.post("/api/tasks/{task_id}/prioritize")
async def prioritize_task(task_id: int, db: AsyncSession = Depends(get_db)):
    task = await db.get(ExpandTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status != "queued":
        raise HTTPException(status_code=400, detail="Only queued tasks can be prioritized")
    max_result = await db.execute(
        select(func.max(ExpandTask.queue_priority)).where(ExpandTask.status == "queued")
    )
    max_priority = max_result.scalar() or 0
    task.queue_priority = max_priority + 1
    task.updated_at = datetime.utcnow()
    await db.commit()
    await _dispatch_next_task()
    return {"task_id": task_id, "status": "queued", "message": "prioritized"}


@app.post("/api/tasks/{task_id}/pause")
async def pause_task(task_id: int, db: AsyncSession = Depends(get_db)):
    task = await db.get(ExpandTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status == "queued":
        task.status = "paused"
        task.updated_at = datetime.utcnow()
        await db.commit()
        return {"task_id": task_id, "status": "paused"}
    if task.status == "running":
        pause_requests.add(task_id)
        task.status = "pausing"
        task.updated_at = datetime.utcnow()
        await db.commit()
        if task_id in cancel_events:
            cancel_events[task_id].set()
        return {"task_id": task_id, "status": "pausing"}
    raise HTTPException(status_code=400, detail="Task cannot be paused in current status")


@app.post("/api/tasks/{task_id}/resume")
async def resume_task(task_id: int, db: AsyncSession = Depends(get_db)):
    task = await db.get(ExpandTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status != "paused":
        raise HTTPException(status_code=400, detail="Only paused tasks can be resumed")
    task.status = "queued"
    task.queue_priority = 0
    task.queued_at = datetime.utcnow()
    task.updated_at = datetime.utcnow()
    await db.commit()
    await _dispatch_next_task()
    return {"task_id": task_id, "status": "queued"}


@app.post("/api/tasks/{task_id}/cancel")
async def cancel_task(task_id: int, db: AsyncSession = Depends(get_db)):
    task = await db.get(ExpandTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status in {"completed", "failed", "cancelled"}:
        return {"task_id": task_id, "status": task.status}
    pause_requests.discard(task_id)
    task.status = "cancelled"
    task.updated_at = datetime.utcnow()
    await db.commit()
    if task_id in cancel_events:
        cancel_events[task_id].set()
    bg = active_tasks.get(task_id)
    if bg and not bg.done():
        bg.cancel()
    await _dispatch_next_task()
    return {"task_id": task_id, "status": "cancelled"}


# ========== 重写段落/章节 路由 ==========

@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/rewrite")
async def rewrite_chapter_endpoint(
    novel_id: int,
    chapter_id: int,
    body: RewriteRequest,
    db: AsyncSession = Depends(get_db),
):
    """重写整个章节（SSE 流式）"""
    stmt = select(Chapter).where(Chapter.id == chapter_id, Chapter.novel_id == novel_id)
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()

    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")

    content = chapter.expanded_content if chapter.expanded_content else chapter.original_content
    model = body.model or config.DEFAULT_MODEL
    instruction = body.instruction or get_default_rewrite_instruction()

    async def generate():
        full_text = ""
        async for chunk in stream_rewrite_paragraph(
            content, "", "", instruction, model
        ):
            full_text += chunk
            yield {
                "event": "chunk",
                "data": json.dumps({"text": chunk}, ensure_ascii=False),
            }

        # 保存结果
        full_text = normalize_output_text(full_text)
        async with async_session() as save_db:
            save_stmt = select(Chapter).where(Chapter.id == chapter_id)
            save_result = await save_db.execute(save_stmt)
            save_chapter = save_result.scalar_one_or_none()
            if save_chapter:
                if _is_no_expansion_notice(full_text, content):
                    save_chapter.status = "skipped"
                    save_chapter.skipped = True
                    save_chapter.error_message = None
                    save_chapter.updated_at = datetime.utcnow()
                    await save_db.commit()
                    yield {
                        "event": "done",
                        "data": json.dumps({"chapter_id": chapter_id, "status": "skipped"}, ensure_ascii=False),
                    }
                    return
                if save_chapter.expanded_content != full_text:
                    save_chapter.expanded_content_prev = save_chapter.expanded_content
                save_chapter.expanded_content = full_text
                save_chapter.status = "completed"
                save_chapter.skipped = False
                save_chapter.error_message = None
                save_chapter.updated_at = datetime.utcnow()
                await save_db.commit()

        yield {
            "event": "done",
            "data": json.dumps({"chapter_id": chapter_id}, ensure_ascii=False),
        }

    return EventSourceResponse(generate())


@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/rewrite-paragraph")
async def rewrite_paragraph_endpoint(
    novel_id: int,
    chapter_id: int,
    body: RewriteParagraphRequest,
    db: AsyncSession = Depends(get_db),
):
    """Deprecated: single-paragraph rewrite is disabled to avoid accidental token waste."""
    raise HTTPException(status_code=410, detail="Single-paragraph rewrite is disabled. Use whole-chapter instruction rewrite.")


@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/insert-prompt")
async def insert_prompt_endpoint(
    novel_id: int,
    chapter_id: int,
    body: InsertPromptRequest,
    db: AsyncSession = Depends(get_db),
):
    """Deprecated: single-paragraph description expansion is disabled."""
    raise HTTPException(status_code=410, detail="Single-paragraph description expansion is disabled. Use whole-chapter instruction rewrite.")


@app.post("/api/novels/{novel_id}/chapters/{chapter_id}/expand-selection")
async def expand_selection_endpoint(
    novel_id: int,
    chapter_id: int,
    body: SelectionExpandRequest,
    db: AsyncSession = Depends(get_db),
):
    """Deprecated: selected-text expansion is disabled."""
    raise HTTPException(status_code=410, detail="Selected-text expansion is disabled. Use whole-chapter instruction rewrite.")


# ========== 导出路由 ==========

@app.get("/api/novels/{novel_id}/export")
async def export_novel(
    novel_id: int,
    format: str = "txt",
    separator_style: str = "classic",
    db: AsyncSession = Depends(get_db),
):
    """导出扩写后的小说为 txt/docx/epub"""
    stmt = select(Novel).options(selectinload(Novel.chapters)).where(Novel.id == novel_id)
    result = await db.execute(stmt)
    novel = result.scalar_one_or_none()

    if not novel:
        raise HTTPException(status_code=404, detail="Novel not found")

    format = (format or "txt").lower()
    if format not in {"txt", "docx", "epub"}:
        raise HTTPException(status_code=400, detail="Unsupported export format")

    safe_title = _sanitize_filename(novel.title)
    filename = f"{safe_title}_expanded.{format}"

    export_dir = os.path.join(config.DATA_DIR, "exports")
    os.makedirs(export_dir, exist_ok=True)
    export_path = os.path.join(export_dir, filename)

    if format == "txt":
        full_text = _build_export_text(novel, separator_style=separator_style)
        with open(export_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        media_type = "text/plain; charset=utf-8"
    elif format == "docx":
        _build_docx(export_path, novel)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        _build_epub(export_path, novel)
        media_type = "application/epub+zip"

    return FileResponse(
        export_path,
        media_type=media_type,
        filename=filename,
    )


# ========== 后台扩写 Worker ==========

async def expand_worker(task_id: int, cancel_event: asyncio.Event, resume_from_index: int = -1):
    """后台扩写 worker

    Args:
        task_id: 任务ID
        cancel_event: 取消信号
        resume_from_index: 从第几个章节开始（-1=从头开始，用于恢复中断的任务）
    """
    logger.info(f"Worker started for task {task_id} (resume_from_index={resume_from_index})")

    try:
        # 获取任务信息（短 session）
        async with async_session() as db:
            task_stmt = select(ExpandTask).where(ExpandTask.id == task_id)
            task_result = await db.execute(task_stmt)
            task = task_result.scalar_one_or_none()

            if not task:
                logger.error(f"Task {task_id} not found")
                return
            if task.status not in {"running", "pausing"}:
                task.status = "running"
                task.updated_at = datetime.utcnow()
                await db.commit()

            novel_id = task.novel_id
            model = task.model
            mode = "one_pass"
            quality = "balanced"
            use_expanded_base = task.use_expanded_as_base or False
            novel_global_summary = ""

            novel_stmt = select(Novel).where(Novel.id == novel_id)
            novel_result = await db.execute(novel_stmt)
            novel = novel_result.scalar_one_or_none()
            if novel and novel.global_summary:
                novel_global_summary = novel.global_summary

            # 获取章节列表
            chapter_ids = None
            if task.chapter_ids_json:
                try:
                    chapter_ids = json.loads(task.chapter_ids_json)
                except json.JSONDecodeError:
                    pass

            if chapter_ids:
                ch_stmt = (
                    select(Chapter)
                    .where(Chapter.novel_id == novel_id)
                    .where(Chapter.id.in_(chapter_ids))
                    .order_by(Chapter.sort_order)
                )
            else:
                ch_stmt = (
                    select(Chapter)
                    .where(Chapter.novel_id == novel_id)
                    .order_by(Chapter.sort_order)
                )

            ch_result = await db.execute(ch_stmt)
            chapters_data = [
                {
                    "id": ch.id,
                    "title": ch.title,
                    "sort_order": ch.sort_order,
                    "weight": max(len(ch.original_content or ""), 1),
                }
                for ch in ch_result.scalars().all()
            ]

        if not chapters_data:
            async with async_session() as db:
                task_obj = await db.get(ExpandTask, task_id)
                if task_obj:
                    task_obj.status = "completed"
                    task_obj.progress = 1.0
                    task_obj.updated_at = datetime.utcnow()
                    await db.commit()
            return

        total_chapters = len(chapters_data)
        total_weight = sum(item["weight"] for item in chapters_data) or total_chapters
        logger.info(f"Task {task_id}: processing {total_chapters} chapters")

        # 更新总章节数
        async with async_session() as db:
            task_obj = await db.get(ExpandTask, task_id)
            if task_obj:
                task_obj.total_chapters = total_chapters
                await db.commit()

        prev_summary = ""
        completed_count = 0
        failed_count = 0
        skipped_count = 0
        failed_ids: List[int] = []
        completed_weight = 0

        for ch_idx, ch_info in enumerate(chapters_data):
            chapter_id = ch_info["id"]
            chapter_title = ch_info["title"]
            chapter_weight = ch_info["weight"]

            # 恢复模式：跳过已处理的章节
            if ch_idx <= resume_from_index:
                # 加载已有摘要作为上下文
                async with async_session() as db:
                    ch = await db.get(Chapter, chapter_id)
                    if ch and ch.summary:
                        prev_summary = ch.summary
                    # 统计已完成/跳过的章节
                    if ch and ch.status == "completed":
                        completed_count += 1
                        completed_weight += chapter_weight
                    elif ch and ch.status == "skipped":
                        skipped_count += 1
                        completed_weight += chapter_weight
                    elif ch and ch.status == "failed":
                        failed_count += 1
                        failed_ids.append(chapter_id)
                logger.info(f"Task {task_id}: Skipping already processed chapter {ch_idx}: {chapter_title}")
                continue

            if ch_idx > 0:
                await _chapter_cooldown(
                    novel_id=novel_id,
                    task_id=task_id,
                    chapter_id=chapter_id,
                    chapter_title=chapter_title,
                    completed_count=completed_count,
                    failed_count=failed_count,
                    skipped_count=skipped_count,
                    total_chapters=total_chapters,
                    completed_weight=completed_weight,
                    total_weight=total_weight,
                    cancel_event=cancel_event,
                )

            # 检查取消
            if cancel_event.is_set():
                logger.info(f"Task {task_id} cancelled at chapter {ch_idx + 1}/{total_chapters}")
                is_pause = task_id in pause_requests
                async with async_session() as db:
                    task_obj = await db.get(ExpandTask, task_id)
                    if task_obj:
                        is_pause = is_pause or task_obj.status == "pausing"
                        task_obj.status = "paused" if is_pause else "cancelled"
                        task_obj.updated_at = datetime.utcnow()
                        await db.commit()
                await broadcast_sse(novel_id, "task_done", {
                    "task_id": task_id,
                    "status": "paused" if is_pause else "cancelled",
                })
                pause_requests.discard(task_id)
                return

            async with async_session() as db:
                task_obj = await db.get(ExpandTask, task_id)
                if task_obj and task_obj.status == "cancelled":
                    logger.info(f"Task {task_id} stopped from database status at chapter {ch_idx + 1}/{total_chapters}")
                    cancel_event.set()
                    await broadcast_sse(novel_id, "task_done", {
                        "task_id": task_id,
                        "status": "cancelled",
                    })
                    return
                if task_obj and task_obj.status in {"paused", "pausing"}:
                    logger.info(f"Task {task_id} paused from database status at chapter {ch_idx + 1}/{total_chapters}")
                    cancel_event.set()
                    if task_obj.status == "pausing":
                        task_obj.status = "paused"
                        task_obj.updated_at = datetime.utcnow()
                        await db.commit()
                    await broadcast_sse(novel_id, "task_done", {
                        "task_id": task_id,
                        "status": "paused",
                    })
                    pause_requests.discard(task_id)
                    return

            logger.info(f"Task {task_id}: Processing chapter {ch_idx + 1}/{total_chapters}: {chapter_title}")

            # 每个章节使用独立的 db session
            try:
                async with async_session() as db:
                    # 重新加载章节数据
                    ch_stmt = select(Chapter).where(Chapter.id == chapter_id)
                    ch_result = await db.execute(ch_stmt)
                    chapter = ch_result.scalar_one_or_none()

                    if not chapter:
                        logger.warning(f"Chapter {chapter_id} not found, skipping")
                        continue

                    # 更新状态
                    chapter.status = "expanding"
                    chapter.progress = 0.0
                    await db.commit()

                    # 广播进度
                    await broadcast_sse(novel_id, "progress", {
                        "task_id": task_id,
                        "chapter_id": chapter.id,
                        "chapter_title": chapter.title,
                        "status": "expanding",
                        "chapter_progress": 0.0,
                        "overall_progress": completed_weight / total_weight,
                        "completed_chapters": completed_count,
                        "failed_chapters": failed_count,
                        "skipped_chapters": skipped_count,
                        "total_chapters": total_chapters,
                    })

                    # 定义进度回调（使用节流更新）
                    async def progress_callback(
                        ch_progress: float,
                        message: str,
                        _chapter_id=chapter.id,
                        _chapter_title=chapter.title,
                        _completed=completed_count,
                        _failed=failed_count,
                        _skipped=skipped_count,
                        _total=total_chapters,
                        _completed_weight=completed_weight,
                        _chapter_weight=chapter_weight,
                    ):
                        overall = (_completed_weight + ch_progress * _chapter_weight) / total_weight
                        await _debounced_progress_update(
                            task_id=task_id,
                            chapter_id=_chapter_id,
                            ch_progress=ch_progress,
                            overall=overall,
                            message=message,
                            novel_id=novel_id,
                            chapter_title=_chapter_title,
                            completed=_completed,
                            failed=_failed,
                            skipped=_skipped,
                            total=_total,
                            current_weight=_chapter_weight,
                        )

                    # 确定输入内容
                    input_content = chapter.original_content
                    previous_expanded_snapshot = chapter.expanded_content
                    if use_expanded_base and chapter.expanded_content:
                        input_content = chapter.expanded_content
                        logger.info(f"继续扩写模式：基于已扩写内容 ({len(chapter.expanded_content)}字)")

                    # 获取下一章短锚点作为衔接上下文。只给很短片段，避免模型把下一章正文复制到本章结尾。
                    next_chapter_opening = ""
                    if ch_idx + 1 < total_chapters:
                        next_ch_id = chapters_data[ch_idx + 1]["id"]
                        try:
                            async with async_session() as next_db:
                                next_ch = await next_db.get(Chapter, next_ch_id)
                                if next_ch and next_ch.original_content:
                                    # 取下一章开头约180字（到句子边界）
                                    opening = next_ch.original_content[:180]
                                    # 在句子边界截断
                                    for sep in ['。', '！', '？', '"', '\n']:
                                        last_pos = opening.rfind(sep)
                                        if last_pos > 40:
                                            opening = opening[:last_pos + 1]
                                            break
                                    next_chapter_opening = f"【{next_ch.title}】\n{opening}"
                        except Exception as e:
                            logger.warning(f"Failed to fetch next chapter opening: {e}")

                    # 中间保存回调：每段扩写完成后保存中间结果到数据库
                    async def segment_save_cb(
                        intermediate_text: str,
                        seg_done: int,
                        seg_total: int,
                        _chapter_id=chapter.id,
                    ):
                        """每段扩写完成后保存中间结果到数据库"""
                        async with async_session() as save_db:
                            ch = await save_db.get(Chapter, _chapter_id)
                            if ch:
                                ch.expanded_content = normalize_output_text(intermediate_text)
                                ch.updated_at = datetime.utcnow()
                                await save_db.commit()
                        logger.info(f"中间保存: {seg_done}/{seg_total} 段")

                    # 执行扩写
                    context_summary = _merge_context_summary(
                        global_summary=novel_global_summary,
                        prev_summary=prev_summary,
                    )

                    ai_refused = False
                    try:
                        expanded = await expand_chapter_one_pass(
                            chapter.title,
                            input_content,
                            context_summary,
                            model=model,
                            quality=quality,
                            progress_callback=progress_callback,
                            next_chapter_opening=next_chapter_opening,
                            skip_if_no_content=config.SKIP_IF_NO_CONTENT,
                            segment_save_callback=segment_save_cb,
                        )
                    except AIRefusalError as e:
                        logger.warning(
                            "Task %s: AI refused chapter %s, falling back to original content: %s",
                            task_id,
                            chapter_title,
                            e,
                        )
                        ai_refused = True
                        expanded = input_content

                    expanded = normalize_output_text(expanded)

                    # If a cancel request arrives during the chapter call, stop before committing results.
                    if cancel_event.is_set():
                        logger.info(f"Task {task_id} cancelled during chapter generation: {chapter_title}")
                        chapter.status = "pending"
                        chapter.error_message = None
                        chapter.updated_at = datetime.utcnow()
                        await db.commit()
                        is_pause = task_id in pause_requests
                        async with async_session() as db2:
                            task_obj = await db2.get(ExpandTask, task_id)
                            if task_obj:
                                is_pause = is_pause or task_obj.status == "pausing"
                                task_obj.status = "paused" if is_pause else "cancelled"
                                task_obj.updated_at = datetime.utcnow()
                                await db2.commit()
                        await broadcast_sse(
                            novel_id,
                            "task_done",
                            {"task_id": task_id, "status": "paused" if is_pause else "cancelled"},
                        )
                        pause_requests.discard(task_id)
                        return

                    # 判断是否被跳过（AI 返回原文、拒绝通知或"无需扩写"通知，均视为跳过）
                    chapter_was_skipped = ai_refused or _is_no_expansion_notice(expanded, input_content)
                    if chapter_was_skipped:
                        chapter.status = "skipped"
                        chapter.skipped = True
                        chapter.progress = 1.0
                        chapter.error_message = None
                        # 若 segment_save_cb 已将通知文本写入 expanded_content，清除它
                        # 确保导出时回落到 original_content，而非通知字符串
                        if ai_refused or (chapter.expanded_content and _is_no_expansion_notice(
                            chapter.expanded_content, input_content
                        )):
                            chapter.expanded_content = None
                        chapter.updated_at = datetime.utcnow()
                        await db.commit()
                        skipped_count += 1
                        completed_weight += chapter_weight
                        logger.info(f"Task {task_id}: Chapter {ch_idx + 1}/{total_chapters} skipped: {chapter_title}")
                    else:
                        # 备份上一版扩写内容（用于撤销）
                        if previous_expanded_snapshot != expanded:
                            chapter.expanded_content_prev = previous_expanded_snapshot
                        chapter.expanded_content = expanded
                        chapter.status = "completed"
                        chapter.skipped = False
                        chapter.progress = 1.0
                        chapter.error_message = None
                        chapter.updated_at = datetime.utcnow()
                        await db.commit()
                        completed_count += 1
                        completed_weight += chapter_weight
                        logger.info(f"Task {task_id}: Chapter {ch_idx + 1}/{total_chapters} completed: {chapter_title}")

                # 更新任务进度（独立 session）
                async with async_session() as db:
                    task_obj = await db.get(ExpandTask, task_id)
                    if task_obj:
                        task_obj.completed_chapters = completed_count
                        task_obj.failed_chapters = failed_count
                        task_obj.skipped_chapters = skipped_count
                        task_obj.last_completed_index = ch_idx
                        task_obj.progress = completed_weight / total_weight
                        task_obj.current_chapter_title = chapter_title
                        task_obj.updated_at = datetime.utcnow()
                        await db.commit()

                # 广播章节完成
                await broadcast_sse(novel_id, "chapter_done", {
                    "chapter_id": chapter_id,
                    "status": "skipped" if chapter_was_skipped else "completed",
                })

                # 生成并保存章节摘要（确保上下文链不断，跳过的章节也生成摘要）
                try:
                    content_for_summary = expanded if expanded != input_content else input_content
                    summary = None
                    if expanded == input_content:
                        async with async_session() as db:
                            ch = await db.get(Chapter, chapter_id)
                            if ch and ch.summary:
                                summary = ch.summary

                    if not summary:
                        if config.CONSERVE_REQUESTS:
                            summary = build_local_chapter_summary(chapter_title, content_for_summary)
                        else:
                            summary = await generate_chapter_summary(content_for_summary, model=model)

                        async with async_session() as db:
                            ch = await db.get(Chapter, chapter_id)
                            if ch:
                                ch.summary = summary
                                await db.commit()
                    prev_summary = summary
                    await _refresh_novel_global_summary(novel_id)
                    async with async_session() as db:
                        novel = await db.get(Novel, novel_id)
                        if novel and novel.global_summary:
                            novel_global_summary = novel.global_summary
                except Exception as e:
                    logger.warning(f"Failed to generate summary: {e}")
                    prev_summary = ""

            except asyncio.CancelledError:
                logger.info(f"Task {task_id} cancelled during chapter {chapter_title}")
                raise
            except ExpansionIntegrityError as e:
                logger.error(f"Integrity check failed for chapter {chapter_title}: {e}")
                failed_count += 1
                failed_ids.append(chapter_id)

                try:
                    async with async_session() as db:
                        fail_ch = await db.get(Chapter, chapter_id)
                        if fail_ch:
                            # 清掉分段中间保存的半成品，避免导出时混入漏剧情/串章文本。
                            fail_ch.expanded_content = previous_expanded_snapshot
                            fail_ch.status = "failed"
                            fail_ch.error_message = str(e)
                            fail_ch.updated_at = datetime.utcnow()
                            await db.commit()
                except Exception as db_err:
                    logger.error(f"Failed to update integrity failure status: {db_err}")

                try:
                    async with async_session() as db:
                        task_obj = await db.get(ExpandTask, task_id)
                        if task_obj:
                            task_obj.failed_chapters = failed_count
                            task_obj.failed_chapter_ids_json = json.dumps(failed_ids)
                            task_obj.last_completed_index = ch_idx
                            task_obj.error_message = str(e)
                            task_obj.updated_at = datetime.utcnow()
                            await db.commit()
                except Exception:
                    pass

                await broadcast_sse(novel_id, "error", {
                    "chapter_id": chapter_id,
                    "error": str(e),
                })
            except Exception as e:
                logger.error(f"Failed to expand chapter {chapter_title}: {e}", exc_info=True)
                # 失败不计入 completed，计入 failed
                failed_count += 1
                failed_ids.append(chapter_id)

                # 标记章节失败
                try:
                    async with async_session() as db:
                        fail_ch = await db.get(Chapter, chapter_id)
                        if fail_ch:
                            fail_ch.status = "failed"
                            fail_ch.error_message = str(e)
                            fail_ch.updated_at = datetime.utcnow()
                            await db.commit()
                except Exception as db_err:
                    logger.error(f"Failed to update chapter status: {db_err}")

                # 更新任务失败信息
                try:
                    async with async_session() as db:
                        task_obj = await db.get(ExpandTask, task_id)
                        if task_obj:
                            task_obj.failed_chapters = failed_count
                            task_obj.failed_chapter_ids_json = json.dumps(failed_ids)
                            task_obj.last_completed_index = ch_idx
                            task_obj.error_message = str(e)
                            task_obj.updated_at = datetime.utcnow()
                            await db.commit()
                except Exception:
                    pass

                await broadcast_sse(novel_id, "error", {
                    "chapter_id": chapter_id,
                    "error": str(e),
                })

                if _is_fatal_api_error(e):
                    logger.error(f"Task {task_id}: fatal API error, stopping batch: {e}")
                    async with async_session() as db:
                        task_obj = await db.get(ExpandTask, task_id)
                        if task_obj:
                            task_obj.status = "failed"
                            task_obj.progress = completed_weight / total_weight
                            task_obj.completed_chapters = completed_count
                            task_obj.failed_chapters = failed_count
                            task_obj.skipped_chapters = skipped_count
                            task_obj.failed_chapter_ids_json = json.dumps(failed_ids) if failed_ids else None
                            task_obj.error_message = str(e)
                            task_obj.updated_at = datetime.utcnow()
                            await db.commit()
                    await broadcast_sse(novel_id, "task_done", {
                        "task_id": task_id,
                        "status": "failed",
                        "error": str(e),
                    })
                    return

                # 非全局 API 错误继续处理下一章

        # 任务完成 — 判断最终状态
        final_status = "completed"
        if failed_count > 0 and completed_count == 0 and skipped_count == 0:
            final_status = "failed"

        async with async_session() as db:
            task_obj = await db.get(ExpandTask, task_id)
            if task_obj:
                task_obj.status = final_status
                task_obj.progress = 1.0
                task_obj.completed_chapters = completed_count
                task_obj.failed_chapters = failed_count
                task_obj.skipped_chapters = skipped_count
                task_obj.failed_chapter_ids_json = json.dumps(failed_ids) if failed_ids else None
                task_obj.updated_at = datetime.utcnow()
                await db.commit()

        await broadcast_sse(novel_id, "task_done", {
            "task_id": task_id,
            "status": final_status,
            "completed_chapters": completed_count,
            "failed_chapters": failed_count,
            "skipped_chapters": skipped_count,
            "total_chapters": total_chapters,
        })

        logger.info(
            f"Task {task_id} {final_status}: "
            f"completed={completed_count}, skipped={skipped_count}, "
            f"failed={failed_count}, total={total_chapters}"
        )

    except asyncio.CancelledError:
        # Task may be cancelled by user-request (cancel_event) or by server shutdown/reload.
        logger.warning(f"Task {task_id} was cancelled by system/user")
        try:
            async with async_session() as err_db:
                err_task = await err_db.get(ExpandTask, task_id)
                if err_task:
                    if err_task.status == "pausing" or task_id in pause_requests:
                        err_task.status = "paused"
                        err_task.error_message = None
                    elif cancel_event.is_set() or err_task.status == "cancelled":
                        err_task.status = "cancelled"
                        err_task.error_message = err_task.error_message or "Task cancelled"
                    else:
                        err_task.status = "interrupted"
                        err_task.error_message = "Task interrupted (server reload or shutdown)"
                    err_task.updated_at = datetime.utcnow()
                    await err_db.commit()

                    await broadcast_sse(err_task.novel_id, "task_done", {
                        "task_id": task_id,
                        "status": err_task.status,
                        "error": err_task.error_message,
                    })
        except Exception:
            pass
    except Exception as e:
        logger.error(f"Worker error for task {task_id}: {e}", exc_info=True)
        try:
            async with async_session() as err_db:
                err_task = await err_db.get(ExpandTask, task_id)
                if err_task:
                    err_task.status = "failed"
                    err_task.error_message = str(e)
                    err_task.updated_at = datetime.utcnow()
                    await err_db.commit()

                    await broadcast_sse(err_task.novel_id, "task_done", {
                        "task_id": task_id,
                        "status": "failed",
                        "error": str(e),
                    })
        except Exception as inner_e:
            logger.error(f"Failed to update task status: {inner_e}")
    finally:
        # 清理
        cancel_events.pop(task_id, None)
        active_tasks.pop(task_id, None)


# ========== 入口点 ==========
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host=config.HOST, port=config.PORT, reload=True)
